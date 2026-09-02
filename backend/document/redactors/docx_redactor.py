"""
DOCX Redactor — FASE 4.7
Substitui dados sensíveis em parágrafos e tabelas de DOCX (python-docx).

Limitação documentada: texto quebrado em múltiplos runs pode não ser
substituído nesta versão MVP. Nesses casos o texto ainda aparece, mas
a versão textual mascarada estará correta.
"""
from __future__ import annotations

import io
import re
from typing import List, Tuple


def _replace_in_text(text: str, pairs: List[Tuple[str, str]]) -> str:
    for value, placeholder in pairs:
        if value in text:
            text = text.replace(value, placeholder)
    return text


def _mask_run_text(run_text: str, pairs: List[Tuple[str, str]]) -> str:
    return _replace_in_text(run_text, pairs)


def redact_docx(
    docx_bytes: bytes,
    findings: list,
    original_text: str,
) -> bytes:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx não instalado")

    doc = Document(io.BytesIO(docx_bytes))

    # Unique (value, placeholder) pairs
    seen: set[str] = set()
    pairs: List[Tuple[str, str]] = []
    for f in findings:
        if f.start is None or f.end is None:
            continue
        value = original_text[f.start:f.end]
        if not value.strip() or value in seen:
            continue
        seen.add(value)
        pairs.append((value, f.placeholder))

    if not pairs:
        return docx_bytes

    def _process_paragraph(para):
        # Try run-level first (preserves formatting)
        for run in para.runs:
            if run.text:
                run.text = _mask_run_text(run.text, pairs)
        # Fallback: check full paragraph text for cross-run matches
        full = para.text
        masked = _replace_in_text(full, pairs)
        if masked != full:
            # If still dirty after run replacement, clear runs and set plain text
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = masked

    for para in doc.paragraphs:
        _process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para)

    # Headers and footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr:
                for para in hdr.paragraphs:
                    _process_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
